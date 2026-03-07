import io
from typing import Union
from fastapi import UploadFile
import logging

# Configure logging
logger = logging.getLogger(__name__)

def parse_resume(file: Union[UploadFile, bytes], filename: str = "") -> str:
    """
    Parse resume from uploaded file and extract text content.
    
    Args:
        file: Uploaded file object or file bytes
        filename: Name of the file (used to determine file type)
        
    Returns:
        Extracted text content from the resume
        
    Raises:
        ValueError: If parsing fails or file type is unsupported
    """
    try:
        logger.info(f"Parsing resume file: {filename}")
        # Read file content
        if isinstance(file, UploadFile):
            # Reset file pointer to beginning
            file.file.seek(0)
            # Read the file content
            content = file.file.read()
            filename = file.filename
            # Reset pointer again for potential future reads
            file.file.seek(0)
        else:
            content = file
        
        # Ensure content is bytes
        if not isinstance(content, bytes):
            raise ValueError(f"Expected bytes content, got {type(content).__name__}")
        
        # Determine file type from extension
        file_ext = filename.lower().split('.')[-1] if filename else ''
        
        # Parse based on file type
        if file_ext == 'pdf':
            logger.debug("Parsing as PDF file")
            text = _parse_pdf(content)
        elif file_ext in ['doc', 'docx']:
            logger.debug("Parsing as DOCX file")
            text = _parse_docx(content)
        elif file_ext == 'txt':
            logger.debug("Parsing as TXT file")
            text = content.decode('utf-8', errors='ignore')
        else:
            logger.warning(f"Unknown file extension '{file_ext}', attempting text decode")
            # Try to decode as text
            try:
                text = content.decode('utf-8', errors='ignore')
            except Exception as decode_error:
                logger.error(f"Failed to decode as text: {str(decode_error)}")
                text = str(content)
        
        logger.info(f"Successfully parsed resume: {len(text)} characters extracted")
        return text
            
    except Exception as e:
        logger.error(f"Resume parsing failed: {str(e)}", exc_info=True)
        raise ValueError(f"Failed to parse resume: {str(e)}")


def _parse_pdf(content: bytes) -> str:
    """Parse PDF file and extract text"""
    try:
        import PyPDF2
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        logger.debug(f"PDF has {len(pdf_reader.pages)} pages")
        
        text = ""
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                logger.debug(f"Extracted {len(page_text)} chars from page {i+1}")
        
        if text.strip():
            return text.strip()
        else:
            # If no text extracted, return a message
            logger.warning("PDF parsed but no text content found - might be scanned/image-based")
            return "PDF parsed but no text content found. This might be a scanned PDF."
            
    except ImportError:
        logger.error("PyPDF2 library not installed")
        raise ValueError("PyPDF2 is not installed. Cannot parse PDF files.")
    except Exception as e:
        logger.error(f"PDF parsing error: {str(e)}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def _parse_docx(content: bytes) -> str:
    """Parse DOCX file and extract text"""
    try:
        import docx
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)
        
        logger.debug(f"DOCX has {len(doc.paragraphs)} paragraphs")
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        if text.strip():
            return text.strip()
        else:
            logger.warning("DOCX parsed but no text content found")
            return "DOCX parsed but no text content found."
            
    except ImportError:
        logger.error("python-docx library not installed")
        raise ValueError("python-docx is not installed. Cannot parse DOCX files.")
    except Exception as e:
        logger.error(f"DOCX parsing error: {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def validate_resume_file(filename: str, file_size: int, max_size: int = 10 * 1024 * 1024) -> tuple[bool, str]:
    """
    Validate resume file before processing.
    
    Args:
        filename: Name of the uploaded file
        file_size: Size of the file in bytes
        max_size: Maximum allowed file size in bytes (default: 10MB)
        
    Returns:
        tuple: (is_valid, error_message)
    """
    # Check file extension
    allowed_extensions = ['.pdf', '.doc', '.docx', '.txt']
    file_ext = '.' + filename.lower().split('.')[-1] if '.' in filename else ''
    
    if file_ext not in allowed_extensions:
        return False, f"Invalid file type. Allowed: {', '.join(allowed_extensions)}"
    
    # Check file size
    if file_size > max_size:
        return False, f"File too large. Maximum size: {max_size // (1024 * 1024)}MB"
    
    return True, ""
